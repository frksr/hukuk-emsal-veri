"""
İhtarname Üretici Sayfası

Türk hukukunda noter ihtarnamesi taslağı üretir.
Kullanıcı taraf bilgilerini, alacak detayını ve türü girer;
LLM ile gerçekçi bir taslak üretilir, .txt ve .docx olarak indirilebilir.
"""
from __future__ import annotations

import io
import sys
from datetime import date, datetime
from decimal import Decimal, InvalidOperation
from pathlib import Path

import streamlit as st

# Proje kökünü PYTHONPATH'e ekle
ROOT = Path(__file__).resolve().parent.parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from services.ihtarname import (  # noqa: E402
    TUR_PROFILLERI,
    YASAL_UYARI,
    ihtarname_olustur,
)

try:
    from llm.provider import status as llm_status
except Exception:  # pragma: no cover
    llm_status = None  # type: ignore


st.set_page_config(
    page_title="İhtarname Üretici",
    page_icon="📨",
    layout="wide",
)


# -----------------------------------------------------------------------------
# Yardımcılar
# -----------------------------------------------------------------------------
def _build_docx(
    metin: str,
    tur_adi: str,
    referanslar: list[str],
    uyari: str,
) -> bytes:
    """python-docx ile basit, başlıklı bir Word dosyası üretir."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, Cm, RGBColor
    except ImportError as e:
        raise RuntimeError(
            "python-docx yüklü değil. `pip install python-docx` ile yükleyin."
        ) from e

    doc = Document()

    # Sayfa kenar boşlukları
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Varsayılan stil
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # HEADER — sayfa üstü
    header = doc.sections[0].header
    h_para = header.paragraphs[0]
    h_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    h_run = h_para.add_run(f"{tur_adi} — Taslak")
    h_run.italic = True
    h_run.font.size = Pt(9)
    h_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # BAŞLIK
    baslik = doc.add_paragraph()
    baslik.alignment = WD_ALIGN_PARAGRAPH.CENTER
    b_run = baslik.add_run("İHTARNAME")
    b_run.bold = True
    b_run.font.size = Pt(16)
    doc.add_paragraph()  # boşluk

    # BODY — İhtarname metni
    # Sistem promptu zaten "İHTARNAME" başlığını üretebilir,
    # bu yüzden ilk satırı tekrar etmemek için filtreliyoruz.
    satirlar = metin.splitlines()
    if satirlar and satirlar[0].strip().upper() in ("İHTARNAME", "IHTARNAME"):
        satirlar = satirlar[1:]
        # Sonraki boş satırı da atla
        while satirlar and not satirlar[0].strip():
            satirlar = satirlar[1:]

    for satir in satirlar:
        p = doc.add_paragraph(satir if satir.strip() else "")
        p.paragraph_format.space_after = Pt(2)
        # Bölüm başlıklarını kalın yap
        ust = satir.strip().upper()
        if ust in (
            "KONU:",
            "AÇIKLAMALAR:",
            "SONUÇ VE TALEP:",
            "İHTAR EDEN (ALACAKLI)",
            "MUHATAP (BORÇLU)",
            "İHTAR EDEN / VEKİLİ",
        ) or ust.startswith("KONU:") or ust.startswith("AÇIKLAMALAR"):
            for run in p.runs:
                run.bold = True

    # Yasal referanslar bölümü
    doc.add_paragraph()
    ref_p = doc.add_paragraph()
    ref_run = ref_p.add_run("Yasal Dayanak: ")
    ref_run.bold = True
    ref_p.add_run(", ".join(referanslar))

    # FOOTER — yasal uyarı
    footer = doc.sections[0].footer
    f_para = footer.paragraphs[0]
    f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = f_para.add_run(uyari)
    f_run.italic = True
    f_run.font.size = Pt(8)
    f_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _decimal_input(label: str, key: str, default: str = "0.00") -> Decimal:
    """Decimal olarak para girdisi."""
    raw = st.text_input(label, value=default, key=key, help="Örn: 12500.50")
    try:
        return Decimal(raw.replace(",", ".").strip() or "0")
    except (InvalidOperation, AttributeError):
        st.warning(f"'{label}' geçersiz sayı, 0 olarak alınıyor.")
        return Decimal("0")


# -----------------------------------------------------------------------------
# SIDEBAR — LLM durumu ve uyarı
# -----------------------------------------------------------------------------
with st.sidebar:
    st.header("LLM Durumu")
    if llm_status is not None:
        try:
            durum = llm_status()
            st.write(f"**Varsayılan:** `{durum['default']}`")
            col1, col2 = st.columns(2)
            with col1:
                st.metric(
                    "Anthropic",
                    "Aktif" if durum["anthropic"] else "Yok",
                )
            with col2:
                st.metric(
                    "Gemini",
                    "Aktif" if durum["gemini"] else "Yok",
                )
            if not (durum["anthropic"] or durum["gemini"]):
                st.error(
                    "Hiçbir LLM sağlayıcısı yapılandırılmamış. "
                    ".env dosyasına ANTHROPIC_API_KEY veya GOOGLE_API_KEY ekleyin."
                )
        except Exception as e:
            st.warning(f"LLM durumu okunamadı: {e}")
    else:
        st.warning("llm/provider.py modülü yüklenemedi.")

    st.divider()
    st.subheader("Yasal Uyarı")
    st.warning(YASAL_UYARI)


# -----------------------------------------------------------------------------
# ANA SAYFA
# -----------------------------------------------------------------------------
st.title("📨 İhtarname Üretici")
st.caption(
    "Türk hukukuna uygun noter ihtarnamesi taslağı üretir. "
    "TBK 117 · İİK 51 · TBK 89 referansları."
)

with st.expander("ℹ️ İhtarname türleri hakkında", expanded=False):
    for kod, profil in TUR_PROFILLERI.items():
        st.markdown(
            f"- **{profil['ad']}** (`{kod}`) — {profil['aciklama']}  \n"
            f"  Süre: {profil['sure_gun']} gün · "
            f"Referanslar: {', '.join(profil['referanslar'])}"
        )

# Form içinde tabs
with st.form("ihtarname_form", clear_on_submit=False):
    tab_taraflar, tab_alacak, tab_tur = st.tabs(
        ["👥 Taraflar", "💰 Alacak Detayı", "📋 Tür & Talepler"]
    )

    # --------------- TARAFLAR ---------------
    with tab_taraflar:
        st.subheader("İhtar Eden (Alacaklı)")
        c1, c2 = st.columns(2)
        with c1:
            alacakli_ad = st.text_input(
                "Alacaklı Ad-Soyad / Unvan *",
                key="alacakli_ad",
                placeholder="Ahmet YILMAZ",
            )
        with c2:
            alacakli_vekil = st.text_input(
                "Alacaklı Vekili (opsiyonel)",
                key="alacakli_vekil",
                placeholder="Av. Mehmet DEMİR",
            )
        alacakli_adres = st.text_area(
            "Alacaklı Adresi *",
            key="alacakli_adres",
            placeholder="Atatürk Cad. No:1 Çankaya/ANKARA",
            height=70,
        )

        st.divider()
        st.subheader("Muhatap (Borçlu)")
        borclu_ad = st.text_input(
            "Borçlu Ad-Soyad / Unvan *",
            key="borclu_ad",
            placeholder="Mehmet KAYA / X Ltd. Şti.",
        )
        borclu_adres = st.text_area(
            "Borçlu Adresi *",
            key="borclu_adres",
            placeholder="İstiklal Cad. No:5 Beyoğlu/İSTANBUL",
            height=70,
        )

    # --------------- ALACAK DETAYI ---------------
    with tab_alacak:
        c1, c2 = st.columns(2)
        with c1:
            anapara = _decimal_input(
                "Anapara (TL) *", key="anapara", default="10000.00"
            )
            faiz_orani = st.number_input(
                "Yıllık Faiz Oranı (%)",
                min_value=0.0,
                max_value=200.0,
                value=24.0,
                step=0.5,
                key="faiz_orani",
            )
        with c2:
            vade_tarihi = st.date_input(
                "Vade Tarihi *",
                value=date.today(),
                key="vade_tarihi",
                format="DD.MM.YYYY",
            )

        neden = st.text_area(
            "Borcun Sebebi *",
            key="neden",
            placeholder=(
                "Örn: 15.03.2025 tarihli mal teslimine ilişkin "
                "faturadan kaynaklanan ödenmemiş satış bedeli"
            ),
            height=80,
        )
        dayanak_belge = st.text_input(
            "Dayanak Belge",
            key="dayanak_belge",
            placeholder="Örn: 2025/123 sayılı fatura, sözleşme, çek/senet",
        )

    # --------------- TÜR & TALEPLER ---------------
    with tab_tur:
        tur_options = list(TUR_PROFILLERI.keys())
        tur_labels = {k: TUR_PROFILLERI[k]["ad"] for k in tur_options}
        tur = st.selectbox(
            "İhtarname Türü *",
            options=tur_options,
            format_func=lambda k: tur_labels[k],
            key="tur",
        )

        profil = TUR_PROFILLERI[tur]
        st.info(
            f"**{profil['ad']}**  \n{profil['aciklama']}  \n"
            f"Önerilen süre: **{profil['sure_gun']} gün** · "
            f"Referanslar: {', '.join(profil['referanslar'])}"
        )

        st.divider()
        st.markdown("**Ek Talepler** (her satıra bir talep)")
        ek_talepler_raw = st.text_area(
            "Ek Talepler",
            key="ek_talepler",
            placeholder=(
                "Faiz, gecikme zammı, icra masrafları, vekalet ücreti, "
                "yargılama giderleri..."
            ),
            height=110,
            label_visibility="collapsed",
        )

    st.divider()
    submitted = st.form_submit_button(
        "📝 İhtarname Üret",
        type="primary",
        use_container_width=True,
    )

# -----------------------------------------------------------------------------
# ÜRETİM
# -----------------------------------------------------------------------------
if submitted:
    # Validasyon
    eksikler = []
    if not alacakli_ad.strip():
        eksikler.append("Alacaklı Ad-Soyad")
    if not alacakli_adres.strip():
        eksikler.append("Alacaklı Adresi")
    if not borclu_ad.strip():
        eksikler.append("Borçlu Ad-Soyad")
    if not borclu_adres.strip():
        eksikler.append("Borçlu Adresi")
    if not neden.strip():
        eksikler.append("Borcun Sebebi")
    if anapara <= 0:
        eksikler.append("Anapara (0'dan büyük olmalı)")

    if eksikler:
        st.error("Lütfen şu zorunlu alanları doldurun: " + ", ".join(eksikler))
        st.stop()

    taraflar = {
        "alacakli_ad": alacakli_ad.strip(),
        "alacakli_adres": alacakli_adres.strip(),
        "borclu_ad": borclu_ad.strip(),
        "borclu_adres": borclu_adres.strip(),
    }
    if alacakli_vekil.strip():
        taraflar["alacakli_vekil"] = alacakli_vekil.strip()

    alacak_detay = {
        "anapara": anapara,
        "faiz_orani": float(faiz_orani),
        "vade_tarihi": vade_tarihi,
        "neden": neden.strip(),
        "dayanak_belge": dayanak_belge.strip() or "—",
    }

    ek_talepler_list = [
        s.strip()
        for s in (ek_talepler_raw or "").splitlines()
        if s.strip()
    ] or None

    with st.spinner("İhtarname taslağı üretiliyor... (10-30 sn sürebilir)"):
        sonuc = ihtarname_olustur(
            tur=tur,
            taraflar=taraflar,
            alacak_detay=alacak_detay,
            ek_talepler=ek_talepler_list,
        )

    if sonuc.get("hata"):
        st.error(f"Hata: {sonuc['hata']}")
        st.stop()

    if not sonuc.get("ihtarname_metni"):
        st.error("İhtarname üretilemedi: boş yanıt döndü.")
        st.stop()

    # Session'a kaydet — sayfa yenilenmeden indirme yapılabilsin
    st.session_state["son_ihtarname"] = sonuc
    st.session_state["son_tur_ad"] = TUR_PROFILLERI[tur]["ad"]

# -----------------------------------------------------------------------------
# SONUÇ GÖRÜNTÜLEME
# -----------------------------------------------------------------------------
if "son_ihtarname" in st.session_state:
    sonuc = st.session_state["son_ihtarname"]
    tur_ad = st.session_state.get("son_tur_ad", "İhtarname")

    st.success(f"✅ {tur_ad} taslağı hazır.")

    # Metadata
    c1, c2, c3 = st.columns(3)
    with c1:
        st.metric("Tahsilat Süresi", f"{sonuc['tahsilat_suresi_gun']} gün")
    with c2:
        st.metric("Yasal Referans", f"{len(sonuc['yasa_referanslari'])} adet")
    with c3:
        st.metric("Tür", sonuc["tur"])

    with st.expander("📚 Yasal Referanslar", expanded=False):
        for ref in sonuc["yasa_referanslari"]:
            st.markdown(f"- {ref}")

    st.warning(f"⚠️ {sonuc['uyari']}")

    st.divider()
    st.subheader("İhtarname Metni")
    st.code(sonuc["ihtarname_metni"], language="text")

    # İndirme butonları
    st.divider()
    st.subheader("📥 İndir")

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    dosya_kok = f"ihtarname_{sonuc['tur']}_{ts}"

    col_txt, col_docx = st.columns(2)
    with col_txt:
        st.download_button(
            label="📄 .txt olarak indir",
            data=sonuc["ihtarname_metni"].encode("utf-8"),
            file_name=f"{dosya_kok}.txt",
            mime="text/plain",
            use_container_width=True,
        )

    with col_docx:
        try:
            docx_bytes = _build_docx(
                metin=sonuc["ihtarname_metni"],
                tur_adi=tur_ad,
                referanslar=sonuc["yasa_referanslari"],
                uyari=sonuc["uyari"],
            )
            st.download_button(
                label="📝 .docx olarak indir",
                data=docx_bytes,
                file_name=f"{dosya_kok}.docx",
                mime=(
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
                use_container_width=True,
            )
        except RuntimeError as e:
            st.error(str(e))
        except Exception as e:
            st.error(f"DOCX üretim hatası: {e}")

    st.divider()
    if st.button("🗑️ Sonucu Temizle", type="secondary"):
        del st.session_state["son_ihtarname"]
        if "son_tur_ad" in st.session_state:
            del st.session_state["son_tur_ad"]
        st.rerun()
