"""
Sözleşme Analiz Servisi

Bir sözleşme dosyasını (PDF/DOCX/TXT) parse eder, maddelerine ayırır
ve LLM ile madde madde risk analizi + genel analiz üretir.

Türk hukukunda TBK, TTK, KVKK çerçevesinde değerlendirme yapar.

UYARI: AI analizi hukuki danışmanlığın yerine geçmez.
"""

from __future__ import annotations

import io
import json
import re
from pathlib import Path
from typing import Any, Iterable

try:
    from llm.provider import generate
except Exception:  # pragma: no cover
    generate = None  # type: ignore


# ---------------------------------------------------------------------------
# Sabitler
# ---------------------------------------------------------------------------

YASAL_UYARI = (
    "AI analizi, hukuki danışmanlığın yerine geçmez. "
    "Sözleşme imzalamadan önce mutlaka bir avukatla görüşün."
)

RISK_SEVIYELERI = ("dusuk", "orta", "yuksek")

# Chunked analiz için: her LLM çağrısına maksimum kaç karakter sözleşme metni gider
CHUNK_KARAKTER = 12000
# Toplam metin bu sınırın altındaysa tek çağrıda analiz et
TEK_CAGRI_SINIRI = 14000

# Maddelerin LLM'e gönderilirken metnini kısaltma sınırı
MADDE_MAX_KARAKTER = 3500

# MALİYET KORUMASI: Tek bir sözleşme analizi için en fazla kaç madde-grubu (LLM
# çağrısı) işlenir. Normal sözleşmeler birkaç chunk eder; bu sınır yalnızca aşırı
# büyük/kötüye kullanım girdilerinde (yüzlerce sayfa) devreye girer ve maliyeti
# sabitler. + 1 genel analiz çağrısı.
MAX_ANALIZ_CHUNK = 16

SOZLESME_TURLERI = {
    "genel": "Genel sözleşme",
    "hizmet": "Hizmet sözleşmesi",
    "satis": "Satış sözleşmesi",
    "kira": "Kira sözleşmesi",
    "gizlilik": "Gizlilik (NDA) sözleşmesi",
    "is": "İş sözleşmesi",
    "distributorluk": "Distribütörlük / Bayilik sözleşmesi",
}

SISTEM_PROMPT = (
    "Sen Türk hukukunda sözleşme uzmanısın. TBK, TTK, KVKK çerçevesinde "
    "sözleşmeleri analiz ediyorsun. Her madde için: 1) sade Türkçe özet, "
    "2) risk seviyesi (düşük/orta/yüksek + gerekçe), 3) iyileştirme önerisi. "
    "Cezai şart, fesih, sorumluluk sınırlama, gizlilik gibi kritik maddelere "
    "özellikle dikkat et. Yanıtını HER ZAMAN geçerli JSON olarak ver, "
    "başka bir açıklama yapma."
)


# ---------------------------------------------------------------------------
# 1) Dosya parse
# ---------------------------------------------------------------------------

def parse_dosya(file_path_or_bytes: Any, ext: str) -> str:
    """Dosyayı (PDF/DOCX/TXT) ham metne çevirir.

    Args:
        file_path_or_bytes: Dosya yolu (str/Path), bytes veya file-like obje
                            (Streamlit UploadedFile gibi).
        ext: Dosya uzantısı (".pdf", ".docx", ".txt" — büyük/küçük harf duyarsız).

    Returns:
        Metnin ham hali. Hata durumunda boş string.
    """
    if not ext:
        return ""
    ext_norm = ext.lower().strip()
    if not ext_norm.startswith("."):
        ext_norm = "." + ext_norm

    # Girdiyi bytes'a normalize et
    veri_bytes: bytes | None = None
    if isinstance(file_path_or_bytes, (str, Path)):
        try:
            veri_bytes = Path(file_path_or_bytes).read_bytes()
        except Exception as e:
            raise RuntimeError(f"Dosya okunamadı: {e}") from e
    elif isinstance(file_path_or_bytes, (bytes, bytearray)):
        veri_bytes = bytes(file_path_or_bytes)
    elif hasattr(file_path_or_bytes, "read"):
        try:
            file_path_or_bytes.seek(0)
        except Exception:
            pass
        veri_bytes = file_path_or_bytes.read()
        if isinstance(veri_bytes, str):
            veri_bytes = veri_bytes.encode("utf-8", errors="ignore")
    else:
        raise TypeError(
            "parse_dosya: file_path_or_bytes str/Path/bytes/file-like olmalı"
        )

    if veri_bytes is None:
        return ""

    if ext_norm == ".pdf":
        return _pdf_oku(veri_bytes)
    if ext_norm == ".docx":
        return _docx_oku(veri_bytes)
    if ext_norm in (".txt", ".text", ".md"):
        return _txt_oku(veri_bytes)

    raise ValueError(
        f"Desteklenmeyen uzantı: {ext_norm}. Sadece .pdf, .docx, .txt desteklenir."
    )


def _pdf_oku(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise RuntimeError(
            "pypdf yüklü değil. `pip install pypdf` ile kurun."
        ) from e
    try:
        reader = PdfReader(io.BytesIO(data))
        parcalar: list[str] = []
        for sayfa in reader.pages:
            try:
                metin = sayfa.extract_text() or ""
            except Exception:
                metin = ""
            if metin:
                parcalar.append(metin)
        return _temizle(("\n".join(parcalar)))
    except Exception as e:
        raise RuntimeError(f"PDF parse hatası: {e}") from e


def _docx_oku(data: bytes) -> str:
    try:
        from docx import Document  # python-docx
    except ImportError as e:
        raise RuntimeError(
            "python-docx yüklü değil. `pip install python-docx` ile kurun."
        ) from e
    try:
        doc = Document(io.BytesIO(data))
        parcalar: list[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parcalar.append(t)
        # Tablolardaki metni de al
        for table in doc.tables:
            for row in table.rows:
                hucreler = [c.text.strip() for c in row.cells if c.text.strip()]
                if hucreler:
                    parcalar.append(" | ".join(hucreler))
        return _temizle("\n".join(parcalar))
    except Exception as e:
        raise RuntimeError(f"DOCX parse hatası: {e}") from e


def _txt_oku(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "windows-1254", "iso-8859-9", "latin-1"):
        try:
            return _temizle(data.decode(enc))
        except UnicodeDecodeError:
            continue
    return _temizle(data.decode("utf-8", errors="ignore"))


def _temizle(metin: str) -> str:
    """Çok sayıdaki boşluk/satırı normalize eder."""
    if not metin:
        return ""
    # Windows satır sonları
    metin = metin.replace("\r\n", "\n").replace("\r", "\n")
    # 3+ boş satırı 2'ye düşür
    metin = re.sub(r"\n{3,}", "\n\n", metin)
    # Satır içi çoklu boşlukları tek boşluk
    metin = re.sub(r"[ \t]+", " ", metin)
    # Satır başı/sonu boşluklar
    metin = "\n".join(s.rstrip() for s in metin.split("\n"))
    return metin.strip()


# ---------------------------------------------------------------------------
# 2) Madde ayırıcı
# ---------------------------------------------------------------------------

# Madde başlığı yakalama regex'leri. Önce dene, başarısız olursa diğerine geç.
# Yakalama grupları: (1) madde numarası
_MADDE_REGEXLERI: list[re.Pattern[str]] = [
    # "MADDE 1", "MADDE 1.", "Madde 12 -", "Madde 3 :"
    re.compile(
        r"(?im)^\s*madde\s+([0-9]{1,3})\s*[\.\-:\)]?\s*",
    ),
    # "1.", "1)", "1 -" satır başında (en fazla 3 haneli numara)
    re.compile(
        r"(?m)^\s*([0-9]{1,3})\s*[\.\)\-]\s+",
    ),
]


def madde_ayir(text: str) -> list[dict]:
    """Sözleşme metnini maddelere ayırır.

    Returns:
        [{"no": "1", "metin": "..."}, ...]
        Eğer madde başlığı yakalanamazsa tek bir "0" madde olarak tüm metni döner.
    """
    if not text or not text.strip():
        return []

    for regex in _MADDE_REGEXLERI:
        maddeler = _maddeleri_yakala(text, regex)
        # En az 2 madde varsa anlamlı kabul et
        if len(maddeler) >= 2:
            return maddeler

    # Hiç madde yakalanmadıysa: tek blok döndür
    return [{"no": "0", "metin": text.strip()}]


def _maddeleri_yakala(text: str, regex: re.Pattern[str]) -> list[dict]:
    eslesmeler = list(regex.finditer(text))
    if not eslesmeler:
        return []

    sonuc: list[dict] = []

    # İlk eşleşmeden öncesi varsa "giris" (başlık/önsöz) olarak ekle
    bas_ofset = eslesmeler[0].start()
    if bas_ofset > 0:
        giris = text[:bas_ofset].strip()
        if giris and len(giris) > 30:  # çok kısa başlıkları yok say
            sonuc.append({"no": "giris", "metin": giris})

    for i, m in enumerate(eslesmeler):
        no = m.group(1).strip()
        baslangic = m.end()
        bitis = eslesmeler[i + 1].start() if i + 1 < len(eslesmeler) else len(text)
        govde = text[baslangic:bitis].strip()
        if not govde:
            continue
        # Madde gövdesi çok kısaysa muhtemelen yanlış eşleşme, atla
        if len(govde) < 10:
            continue
        sonuc.append({"no": no, "metin": govde})

    return sonuc


# ---------------------------------------------------------------------------
# 3) LLM ile analiz
# ---------------------------------------------------------------------------

def _llm_cagir(system: str, user: str, max_tokens: int = 2000) -> str:
    """LLM'i çağırır, ham string döner."""
    if generate is None:
        raise RuntimeError("LLM provider yüklenemedi (llm/provider.py).")
    try:
        return generate(
            system=system,
            user=user,
            max_tokens=max_tokens,
            temperature=0.2,
        )
    except TypeError:
        return generate(system=system, user=user, max_tokens=max_tokens)


_JSON_KOD_BLOK_RE = re.compile(r"```(?:json)?\s*(\{[\s\S]*?\})\s*```", re.IGNORECASE)


def _json_parse_guvenli(ham: str) -> dict | None:
    """LLM çıktısından JSON'u kurtarmaya çalışır."""
    if not ham:
        return None
    ham = ham.strip()
    # Önce kod bloğu içinde JSON ara
    m = _JSON_KOD_BLOK_RE.search(ham)
    if m:
        try:
            return json.loads(m.group(1))
        except Exception:
            pass
    # Direkt JSON dene
    try:
        return json.loads(ham)
    except Exception:
        pass
    # İlk { ve son } arasını dene
    bas = ham.find("{")
    son = ham.rfind("}")
    if bas != -1 and son != -1 and son > bas:
        try:
            return json.loads(ham[bas : son + 1])
        except Exception:
            return None
    return None


def _risk_normalize(deger: Any) -> str:
    if not isinstance(deger, str):
        return "orta"
    d = deger.strip().lower()
    d = (
        d.replace("ü", "u")
        .replace("ş", "s")
        .replace("ı", "i")
        .replace("ö", "o")
        .replace("ç", "c")
        .replace("ğ", "g")
    )
    if "yuksek" in d or "high" in d:
        return "yuksek"
    if "dusuk" in d or "low" in d:
        return "dusuk"
    return "orta"


def _madde_payload(maddeler: list[dict]) -> list[dict]:
    """LLM'e gönderilecek kompakt madde listesi (uzun maddeler kısaltılır)."""
    out: list[dict] = []
    for m in maddeler:
        metin = m.get("metin", "")
        if len(metin) > MADDE_MAX_KARAKTER:
            metin = metin[:MADDE_MAX_KARAKTER] + " [...kısaltıldı...]"
        out.append({"no": str(m.get("no", "?")), "metin": metin})
    return out


def _chunk_madde_listesi(
    maddeler: list[dict], chunk_char: int = CHUNK_KARAKTER
) -> Iterable[list[dict]]:
    """Maddeleri toplam karakter sınırını aşmayacak gruplara böler."""
    grup: list[dict] = []
    grup_uzunluk = 0
    for m in maddeler:
        m_uzunluk = len(m.get("metin", "")) + 50  # JSON overhead
        if grup and grup_uzunluk + m_uzunluk > chunk_char:
            yield grup
            grup = []
            grup_uzunluk = 0
        grup.append(m)
        grup_uzunluk += m_uzunluk
    if grup:
        yield grup


def _madde_grubu_analiz(
    grup: list[dict], sozlesme_turu: str
) -> list[dict]:
    """Bir madde grubunu LLM'e gönderip madde madde analiz alır."""
    payload = _madde_payload(grup)
    user_prompt = (
        f"Sözleşme türü: {SOZLESME_TURLERI.get(sozlesme_turu, sozlesme_turu)}\n\n"
        "Aşağıdaki sözleşme maddelerini madde madde analiz et. "
        "Her madde için: kısa Türkçe özet, risk_seviye (dusuk/orta/yuksek), "
        "gerekçeli iyileştirme önerisi ve atıf yapılan/ilgili TBK/TTK/KVKK "
        "kanun maddeleri.\n\n"
        "SADECE aşağıdaki JSON şemasında yanıt ver:\n"
        "{\n"
        '  "maddeler": [\n'
        '    {"no": "1", "ozet": "...", "risk_seviye": "dusuk|orta|yuksek", '
        '"oneri": "...", "ilgili_kanunlar": ["TBK m. 138", "..."]}\n'
        "  ]\n"
        "}\n\n"
        f"MADDELER (JSON):\n{json.dumps(payload, ensure_ascii=False)}"
    )

    try:
        ham = _llm_cagir(SISTEM_PROMPT, user_prompt, max_tokens=3500)
    except Exception as e:
        # LLM hata verdiyse fallback olarak her maddeyi "analiz edilemedi" yap
        return [
            {
                "no": m["no"],
                "ozet": "(analiz edilemedi)",
                "risk_seviye": "orta",
                "oneri": f"LLM hatası: {e}",
                "ilgili_kanunlar": [],
            }
            for m in grup
        ]

    parsed = _json_parse_guvenli(ham)
    sonuclar: list[dict] = []
    if parsed and isinstance(parsed.get("maddeler"), list):
        # no -> analiz map'i kur, eksik kalanlar için fallback
        analiz_map = {}
        for item in parsed["maddeler"]:
            if not isinstance(item, dict):
                continue
            no = str(item.get("no", "")).strip()
            if not no:
                continue
            analiz_map[no] = item

        for m in grup:
            no = str(m["no"])
            it = analiz_map.get(no, {})
            kanunlar = it.get("ilgili_kanunlar", []) or []
            if isinstance(kanunlar, str):
                kanunlar = [kanunlar]
            sonuclar.append(
                {
                    "no": no,
                    "ozet": str(it.get("ozet", "")).strip()
                    or "(özet üretilemedi)",
                    "risk_seviye": _risk_normalize(it.get("risk_seviye")),
                    "oneri": str(it.get("oneri", "")).strip(),
                    "ilgili_kanunlar": [str(k).strip() for k in kanunlar if k],
                }
            )
    else:
        # JSON parse başarısız — fallback
        for m in grup:
            sonuclar.append(
                {
                    "no": str(m["no"]),
                    "ozet": "(LLM çıktısı parse edilemedi)",
                    "risk_seviye": "orta",
                    "oneri": "Tekrar deneyin veya manuel inceleyin.",
                    "ilgili_kanunlar": [],
                }
            )

    return sonuclar


def _genel_analiz(
    sozlesme_metni: str, sozlesme_turu: str
) -> dict:
    """Sözleşmenin genel özet/taraf/konu/eksik maddeler analizi."""
    # Çok uzun ise kısalt (baş + son)
    metin = sozlesme_metni
    if len(metin) > TEK_CAGRI_SINIRI:
        yarim = TEK_CAGRI_SINIRI // 2
        metin = (
            metin[:yarim]
            + "\n\n[...orta kısım atlandı...]\n\n"
            + metin[-yarim:]
        )

    user_prompt = (
        f"Sözleşme türü: {SOZLESME_TURLERI.get(sozlesme_turu, sozlesme_turu)}\n\n"
        "Aşağıdaki sözleşmenin GENEL analizini yap. "
        "Yanıtını SADECE şu JSON şemasında ver:\n"
        "{\n"
        '  "genel_ozet": "2-4 cümle, sözleşmenin amacı",\n'
        '  "taraflar": ["Taraf A", "Taraf B"],\n'
        '  "ana_konu": "kısa cümle",\n'
        '  "sure_ve_fesih": "süre ve fesih hükümlerinin özeti",\n'
        '  "eksik_maddeler": ["Eklenmesi önerilen madde 1", "..."],\n'
        '  "ana_riskler": ["Risk 1", "..."],\n'
        '  "uyari": "Tarafları en çok bağlayan/risk taşıyan nokta"\n'
        "}\n\n"
        "Eksik maddeler kısmında: bu tür sözleşmelerde olması gerekip "
        "metinde olmayan hükümleri (örn. mücbir sebep, uyuşmazlık çözümü, "
        "KVKK, gizlilik, fesih bildirimi) listele.\n\n"
        f"--- SÖZLEŞME BAŞLANGIÇ ---\n{metin}\n--- SÖZLEŞME SON ---"
    )

    try:
        ham = _llm_cagir(SISTEM_PROMPT, user_prompt, max_tokens=2000)
    except Exception as e:
        return {
            "genel_ozet": f"(LLM hatası: {e})",
            "taraflar": [],
            "ana_konu": "",
            "sure_ve_fesih": "",
            "eksik_maddeler": [],
            "ana_riskler": [],
            "uyari": "",
        }

    parsed = _json_parse_guvenli(ham) or {}

    def _liste(d: Any) -> list[str]:
        if isinstance(d, list):
            return [str(x).strip() for x in d if str(x).strip()]
        if isinstance(d, str):
            return [d.strip()] if d.strip() else []
        return []

    return {
        "genel_ozet": str(parsed.get("genel_ozet", "")).strip(),
        "taraflar": _liste(parsed.get("taraflar")),
        "ana_konu": str(parsed.get("ana_konu", "")).strip(),
        "sure_ve_fesih": str(parsed.get("sure_ve_fesih", "")).strip(),
        "eksik_maddeler": _liste(parsed.get("eksik_maddeler")),
        "ana_riskler": _liste(parsed.get("ana_riskler")),
        "uyari": str(parsed.get("uyari", "")).strip(),
    }


def _toplam_risk_skoru(madde_analizleri: list[dict]) -> int:
    """0-100 arası bir risk skoru hesaplar.

    Ağırlık: dusuk=10, orta=40, yuksek=85.
    Sonuç maddelerin ağırlıklı ortalamasıdır.
    """
    if not madde_analizleri:
        return 0
    agirlik = {"dusuk": 10, "orta": 40, "yuksek": 85}
    toplam = 0
    say = 0
    for m in madde_analizleri:
        seviye = _risk_normalize(m.get("risk_seviye"))
        toplam += agirlik.get(seviye, 40)
        say += 1
    if say == 0:
        return 0
    return min(100, max(0, round(toplam / say)))


def analiz_et(
    sozlesme_metni: str, sozlesme_turu: str = "genel"
) -> dict:
    """Sözleşmeyi madde madde + genel olarak analiz eder.

    Args:
        sozlesme_metni: Ham sözleşme metni (parse_dosya çıktısı).
        sozlesme_turu: hizmet, satis, kira, gizlilik, is, distributorluk, genel.

    Returns:
        {
          "genel_ozet": str,
          "taraflar": [str],
          "ana_konu": str,
          "sure_ve_fesih": str,
          "eksik_maddeler": [str],
          "madde_analizleri": [
              {"no": str, "ozet": str, "risk_seviye": "...",
               "oneri": str, "ilgili_kanunlar": [str]}
          ],
          "toplam_risk_skoru": int (0-100),
          "uyari": str,
          ...
        }
    """
    if not sozlesme_metni or not sozlesme_metni.strip():
        return {
            "genel_ozet": "",
            "taraflar": [],
            "ana_konu": "",
            "sure_ve_fesih": "",
            "eksik_maddeler": [],
            "ana_riskler": [],
            "madde_analizleri": [],
            "toplam_risk_skoru": 0,
            "uyari": "Sözleşme metni boş.",
            "sozlesme_turu": sozlesme_turu,
            "yasal_uyari": YASAL_UYARI,
            "hata": "Sözleşme metni boş.",
        }

    if generate is None:
        return {
            "genel_ozet": "",
            "taraflar": [],
            "ana_konu": "",
            "sure_ve_fesih": "",
            "eksik_maddeler": [],
            "ana_riskler": [],
            "madde_analizleri": [],
            "toplam_risk_skoru": 0,
            "uyari": "LLM provider yüklenemedi.",
            "sozlesme_turu": sozlesme_turu,
            "yasal_uyari": YASAL_UYARI,
            "hata": "LLM provider yüklenemedi (llm/provider.py).",
        }

    if sozlesme_turu not in SOZLESME_TURLERI:
        sozlesme_turu = "genel"

    # 1) Maddeleri ayır
    maddeler = madde_ayir(sozlesme_metni)

    # 2) Madde analizini chunked LLM çağrılarıyla yap
    madde_analizleri: list[dict] = []
    if maddeler:
        for i, grup in enumerate(_chunk_madde_listesi(maddeler)):
            # Maliyet koruması: aşırı büyük sözleşmelerde LLM çağrı sayısını sınırla.
            if i >= MAX_ANALIZ_CHUNK:
                break
            madde_analizleri.extend(_madde_grubu_analiz(grup, sozlesme_turu))

    # 3) Genel analiz
    genel = _genel_analiz(sozlesme_metni, sozlesme_turu)

    # 4) Toplam risk skoru
    risk_skoru = _toplam_risk_skoru(madde_analizleri)

    return {
        "genel_ozet": genel["genel_ozet"],
        "taraflar": genel["taraflar"],
        "ana_konu": genel["ana_konu"],
        "sure_ve_fesih": genel["sure_ve_fesih"],
        "eksik_maddeler": genel["eksik_maddeler"],
        "ana_riskler": genel["ana_riskler"],
        "madde_analizleri": madde_analizleri,
        "toplam_risk_skoru": risk_skoru,
        "uyari": genel["uyari"] or YASAL_UYARI,
        "sozlesme_turu": sozlesme_turu,
        "madde_sayisi": len(maddeler),
        "karakter_sayisi": len(sozlesme_metni),
        "yasal_uyari": YASAL_UYARI,
    }


# ---------------------------------------------------------------------------
# Yardımcı: analiz sonucundan .docx / .pdf rapor üret
# ---------------------------------------------------------------------------

def rapor_docx(analiz: dict, dosya_adi: str = "sozlesme_analizi.docx") -> bytes:
    """Analiz sonucundan .docx rapor üretir, bytes döner."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError as e:
        raise RuntimeError("python-docx kurulu değil.") from e

    renk_map = {
        "dusuk": RGBColor(0x16, 0x7A, 0x3A),  # yeşil
        "orta": RGBColor(0xB8, 0x86, 0x0B),   # sarı/turuncu
        "yuksek": RGBColor(0xB0, 0x1B, 0x1B),  # kırmızı
    }
    seviye_etiket = {"dusuk": "DÜŞÜK", "orta": "ORTA", "yuksek": "YÜKSEK"}

    doc = Document()
    doc.add_heading("Sözleşme Analiz Raporu", level=0)

    p = doc.add_paragraph()
    p.add_run(
        f"Sözleşme türü: {SOZLESME_TURLERI.get(analiz.get('sozlesme_turu','genel'), 'Genel')}\n"
        f"Madde sayısı: {analiz.get('madde_sayisi', 0)}\n"
        f"Toplam risk skoru: {analiz.get('toplam_risk_skoru', 0)} / 100"
    )

    doc.add_heading("Genel Özet", level=1)
    doc.add_paragraph(analiz.get("genel_ozet", "") or "—")

    if analiz.get("taraflar"):
        doc.add_heading("Taraflar", level=2)
        for t in analiz["taraflar"]:
            doc.add_paragraph(t, style="List Bullet")

    if analiz.get("ana_konu"):
        doc.add_heading("Ana Konu", level=2)
        doc.add_paragraph(analiz["ana_konu"])

    if analiz.get("sure_ve_fesih"):
        doc.add_heading("Süre ve Fesih", level=2)
        doc.add_paragraph(analiz["sure_ve_fesih"])

    if analiz.get("ana_riskler"):
        doc.add_heading("Ana Riskler", level=2)
        for r in analiz["ana_riskler"]:
            doc.add_paragraph(r, style="List Bullet")

    doc.add_heading("Madde Madde Analiz", level=1)
    for m in analiz.get("madde_analizleri", []):
        no = m.get("no", "?")
        seviye = _risk_normalize(m.get("risk_seviye"))
        h = doc.add_paragraph()
        run_no = h.add_run(f"Madde {no} — ")
        run_no.bold = True
        run_risk = h.add_run(f"Risk: {seviye_etiket.get(seviye, seviye.upper())}")
        run_risk.bold = True
        run_risk.font.color.rgb = renk_map.get(seviye, RGBColor(0, 0, 0))

        if m.get("ozet"):
            doc.add_paragraph(m["ozet"])
        if m.get("oneri"):
            p = doc.add_paragraph()
            p.add_run("Öneri: ").bold = True
            p.add_run(m["oneri"])
        if m.get("ilgili_kanunlar"):
            p = doc.add_paragraph()
            p.add_run("İlgili Kanunlar: ").bold = True
            p.add_run(", ".join(m["ilgili_kanunlar"]))

    if analiz.get("eksik_maddeler"):
        doc.add_heading("Eksik / Eklenmesi Önerilen Maddeler", level=1)
        for e in analiz["eksik_maddeler"]:
            doc.add_paragraph(e, style="List Bullet")

    doc.add_heading("Yasal Uyarı", level=2)
    p = doc.add_paragraph(analiz.get("yasal_uyari", YASAL_UYARI))
    for r in p.runs:
        r.font.size = Pt(9)
        r.italic = True

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def rapor_pdf(analiz: dict) -> bytes:
    """Analiz sonucundan basit bir PDF rapor üretir (reportlab varsa)."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem,
        )
    except ImportError as e:
        raise RuntimeError(
            "reportlab kurulu değil. `pip install reportlab` ile kurun."
        ) from e

    bio = io.BytesIO()
    doc = SimpleDocTemplate(
        bio, pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
    )
    styles = getSampleStyleSheet()
    h1 = styles["Heading1"]
    h2 = styles["Heading2"]
    h3 = styles["Heading3"]
    body = styles["BodyText"]

    risk_renk = {"dusuk": "green", "orta": "orange", "yuksek": "red"}
    seviye_etiket = {"dusuk": "DÜŞÜK", "orta": "ORTA", "yuksek": "YÜKSEK"}

    story = []
    story.append(Paragraph("Sözleşme Analiz Raporu", h1))
    story.append(Paragraph(
        f"Sözleşme türü: {SOZLESME_TURLERI.get(analiz.get('sozlesme_turu','genel'), 'Genel')}<br/>"
        f"Madde sayısı: {analiz.get('madde_sayisi', 0)}<br/>"
        f"Toplam risk skoru: <b>{analiz.get('toplam_risk_skoru', 0)} / 100</b>",
        body,
    ))
    story.append(Spacer(1, 0.4 * cm))

    story.append(Paragraph("Genel Özet", h2))
    story.append(Paragraph(analiz.get("genel_ozet", "") or "—", body))

    if analiz.get("taraflar"):
        story.append(Paragraph("Taraflar", h3))
        story.append(ListFlowable(
            [ListItem(Paragraph(t, body)) for t in analiz["taraflar"]],
            bulletType="bullet",
        ))

    if analiz.get("ana_konu"):
        story.append(Paragraph("Ana Konu", h3))
        story.append(Paragraph(analiz["ana_konu"], body))

    if analiz.get("sure_ve_fesih"):
        story.append(Paragraph("Süre ve Fesih", h3))
        story.append(Paragraph(analiz["sure_ve_fesih"], body))

    if analiz.get("ana_riskler"):
        story.append(Paragraph("Ana Riskler", h3))
        story.append(ListFlowable(
            [ListItem(Paragraph(r, body)) for r in analiz["ana_riskler"]],
            bulletType="bullet",
        ))

    story.append(Paragraph("Madde Madde Analiz", h2))
    for m in analiz.get("madde_analizleri", []):
        no = m.get("no", "?")
        seviye = _risk_normalize(m.get("risk_seviye"))
        renk = risk_renk.get(seviye, "black")
        story.append(Paragraph(
            f"<b>Madde {no}</b> — "
            f"<font color='{renk}'><b>Risk: {seviye_etiket.get(seviye, seviye.upper())}</b></font>",
            body,
        ))
        if m.get("ozet"):
            story.append(Paragraph(m["ozet"], body))
        if m.get("oneri"):
            story.append(Paragraph(f"<b>Öneri:</b> {m['oneri']}", body))
        if m.get("ilgili_kanunlar"):
            story.append(Paragraph(
                f"<b>İlgili Kanunlar:</b> {', '.join(m['ilgili_kanunlar'])}",
                body,
            ))
        story.append(Spacer(1, 0.2 * cm))

    if analiz.get("eksik_maddeler"):
        story.append(Paragraph("Eksik / Eklenmesi Önerilen Maddeler", h2))
        story.append(ListFlowable(
            [ListItem(Paragraph(e, body)) for e in analiz["eksik_maddeler"]],
            bulletType="bullet",
        ))

    story.append(Spacer(1, 0.5 * cm))
    italic = ParagraphStyle("italic", parent=body, fontSize=9, italic=True, textColor="#666666")
    story.append(Paragraph(analiz.get("yasal_uyari", YASAL_UYARI), italic))

    doc.build(story)
    return bio.getvalue()
