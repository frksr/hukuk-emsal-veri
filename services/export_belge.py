"""Belge export servisi — AI çıktılarını .docx ve .udf (UYAP) formatına çevirir.

Avukat iş akışı: taslağı Word'de düzenle → UYAP'a .udf olarak yükle.
.txt indirme yerine bu iki format sunulur.

UDF formatı: UYAP Doküman Editörü'nün dosya biçimi — içinde `content.xml`
bulunan bir ZIP arşividir. content.xml, CDATA içinde düz metni ve
<elements> altında paragraf offset'lerini taşır. Bu üretici, UYAP editörünün
açabildiği minimal-geçerli bir belge oluşturur.
"""
from __future__ import annotations

import io
import zipfile
from xml.sax.saxutils import escape


# -----------------------------------------------------------------------------
# DOCX
# -----------------------------------------------------------------------------

def belge_docx(metin: str, baslik: str | None = None) -> bytes:
    """Düz metni basit ama düzgün biçimli bir .docx'e çevirir.

    - 12pt Times New Roman (dilekçe geleneği), 1.15 satır aralığı
    - Boş satırlar paragraf ayracı kabul edilir
    - TAMAMEN BÜYÜK YAZILMIŞ satırlar ortalanır (MAHKEME başlıkları)
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    if baslik:
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h.add_run(baslik.strip())
        run.bold = True
        run.font.size = Pt(13)
        doc.add_paragraph()

    for satir in (metin or "").splitlines():
        p = doc.add_paragraph()
        text = satir.rstrip()
        # Mahkeme başlığı gibi tamamı büyük satırları ortala
        harfler = [c for c in text if c.isalpha()]
        if harfler and all(c.isupper() for c in harfler) and len(text) < 80:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
        else:
            p.add_run(text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# -----------------------------------------------------------------------------
# UDF (UYAP Doküman Editörü)
# -----------------------------------------------------------------------------

_UDF_TEMPLATE = """<?xml version="1.0" encoding="UTF-8" ?>
<template format_id="1.8">
<content><![CDATA[{content}]]></content>
<properties><pageFormat mediaSizeName="1" leftMargin="70.875" rightMargin="70.875" topMargin="70.875" bottomMargin="70.875" paperOrientation="1" headerFOffset="20.0" footerFOffset="20.0" /></properties>
<elements resolver="hvl-default">
{elements}
</elements>
<styles><style name="default" description="Geçerli" family="Dialog" size="12" bold="false" italic="false" foreground="-13421773" FONT_ATTRIBUTE_KEY="javax.swing.plaf.FontUIResource[family=Dialog,name=Dialog,style=plain,size=12]" /><style name="hvl-default" family="Times New Roman" size="12" description="Gövde" /></styles>
</template>
"""


def belge_udf(metin: str, baslik: str | None = None) -> bytes:
    """Düz metni UYAP .udf belgesine çevirir.

    content.xml yapısı: tüm metin CDATA'da tek blok; <elements> altında her
    paragraf için startOffset/length kaydı tutulur.
    """
    parcalar: list[str] = []
    if baslik:
        parcalar.append(baslik.strip())
        parcalar.append("")
    parcalar.extend((metin or "").splitlines())

    # CDATA içeriği ve paragraf offset'leri
    content_parts: list[str] = []
    elements: list[str] = []
    offset = 0
    for i, satir in enumerate(parcalar):
        text = satir.rstrip("\r")
        # CDATA güvenliği: ]]> dizisini boz
        text = text.replace("]]>", "]] >")
        line = text + "\n"
        content_parts.append(line)
        alignment = ""
        harfler = [c for c in text if c.isalpha()]
        if harfler and all(c.isupper() for c in harfler) and len(text) < 80:
            alignment = ' Alignment="1"'  # ortala (başlıklar)
        elements.append(
            f'<paragraph{alignment}><content startOffset="{offset}" '
            f'length="{len(line)}" /></paragraph>'
        )
        offset += len(line)

    xml = _UDF_TEMPLATE.format(
        content="".join(content_parts),
        elements="\n".join(elements),
    )

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("content.xml", xml.encode("utf-8"))
    return buf.getvalue()


def guvenli_dosya_adi(ad: str, uzanti: str) -> str:
    """İndirilen dosya adı için güvenli slug üretir."""
    import re
    import unicodedata

    s = unicodedata.normalize("NFKD", ad or "belge")
    s = s.encode("ascii", "ignore").decode("ascii")
    s = re.sub(r"[^A-Za-z0-9_-]+", "-", s).strip("-").lower() or "belge"
    return f"{s[:60]}.{uzanti}"


__all__ = ["belge_docx", "belge_udf", "guvenli_dosya_adi"]
