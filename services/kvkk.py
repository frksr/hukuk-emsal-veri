"""KVKK Uyum Kontrol Listesi Servisi.

Sektör + işlenen veri türleri seçildiğinde KVKK (6698 sayılı Kişisel Verilerin
Korunması Kanunu) uyumluluğu için checklist üretir.

Yaklaşım:
  - TEMEL_CHECKLIST: Her durumda gereken statik maddeler (kanun referanslı).
  - SEKTOR_EK_MADDELER: Sektöre özel ek statik maddeler.
  - VERI_TURU_EK_MADDELER: İşlenen veri türüne özel ek statik maddeler.
  - LLM (opsiyonel): Sektöre özel ek maddeler ve sektörel notlar üretir.

UYARI: Bu rehber YAKLAŞIK bilgi içindir ve profesyonel KVKK danışmanlığının
yerine geçmez. Somut uyumluluk için bir hukuk müşaviri / KVKK uzmanına
danışılması gereklidir.
"""

from __future__ import annotations

import json
import re
from typing import Dict, List, Optional

try:
    from llm.provider import generate as llm_generate
    from llm.provider import is_available as llm_available
except ImportError:  # pragma: no cover - fallback for stand-alone import
    llm_generate = None  # type: ignore[assignment]
    llm_available = None  # type: ignore[assignment]


# --------------------------------------------------------------------------- #
# Statik kategoriler ve etiketler
# --------------------------------------------------------------------------- #

KATEGORI_ETIKETLERI: Dict[str, str] = {
    "bilgilendirme": "Bilgilendirme ve Şeffaflık",
    "rıza": "Rıza ve Hukuki Sebep",
    "envanter": "Veri Envanteri ve VERBİS",
    "saklama": "Saklama ve İmha",
    "ihlal": "Veri İhlali Yönetimi",
    "yönetim": "Yönetişim ve Sorumluluk",
    "eğitim": "Farkındalık ve Eğitim",
    "teknik": "Teknik Tedbirler",
    "idari": "İdari Tedbirler",
    "transfer": "Yurt Dışı / 3. Taraf Aktarım",
    "haklar": "İlgili Kişi Hakları",
    "ozel_nitelikli": "Özel Nitelikli Veri",
    "sektorel": "Sektörel Yükümlülükler",
}


SEKTOR_ETIKETLERI: Dict[str, str] = {
    "saglik": "Sağlık",
    "fintech": "Finans / Fintech",
    "egitim": "Eğitim",
    "e_ticaret": "E-Ticaret",
    "imalat": "İmalat / Sanayi",
    "kamu": "Kamu Kurumu",
    "telekom": "Telekomünikasyon",
    "insan_kaynaklari": "İnsan Kaynakları / İK",
    "hukuk_burosu": "Hukuk Bürosu",
    "diger": "Diğer / Genel",
}


VERI_TURU_ETIKETLERI: Dict[str, str] = {
    "kisisel": "Kişisel veri (ad, iletişim, kimlik)",
    "ozel_nitelikli": "Özel nitelikli veri (sağlık, inanç, biyometrik vb.)",
    "finansal": "Finansal veri (IBAN, kart, gelir bilgisi)",
    "saglik": "Sağlık verisi",
    "cocuk": "Çocuk verisi (18 yaş altı)",
    "calisan": "Çalışan / İK verisi",
    "musteri": "Müşteri verisi",
    "konum": "Konum verisi",
    "biyometrik": "Biyometrik veri",
    "iletisim_kaydi": "İletişim / çağrı kaydı",
    "kamera": "Kamera / görüntü kaydı (CCTV)",
    "cerez": "Çerez / dijital iz",
}


# --------------------------------------------------------------------------- #
# Temel Checklist (her durumda lazım)
# --------------------------------------------------------------------------- #

TEMEL_CHECKLIST: List[Dict[str, object]] = [
    {
        "no": 1,
        "madde": "Aydınlatma metni hazırlanmış mı? (KVKK 10)",
        "kategori": "bilgilendirme",
        "oncelik": "yuksek",
    },
    {
        "no": 2,
        "madde": "Açık rıza alınıyor mu, geri çekilebilir mi? (KVKK 5/1)",
        "kategori": "rıza",
        "oncelik": "yuksek",
    },
    {
        "no": 3,
        "madde": "Veri envanteri çıkarılmış mı? (Veri sorumlusu sicili / VERBİS)",
        "kategori": "envanter",
        "oncelik": "yuksek",
    },
    {
        "no": 4,
        "madde": "Veri saklama ve imha politikası var mı? (Yön. Md. 7)",
        "kategori": "saklama",
        "oncelik": "yuksek",
    },
    {
        "no": 5,
        "madde": "Veri ihlali bildirim prosedürü hazır mı? (KVKK 12/5 — 72 saat)",
        "kategori": "ihlal",
        "oncelik": "yuksek",
    },
    {
        "no": 6,
        "madde": "Veri sorumlusu temsilcisi atandı mı? (KVKK 11)",
        "kategori": "yönetim",
        "oncelik": "orta",
    },
    {
        "no": 7,
        "madde": "Çalışan eğitimi yapılıyor mu? (yıllık asgari bir tekrar)",
        "kategori": "eğitim",
        "oncelik": "orta",
    },
    {
        "no": 8,
        "madde": "Teknik tedbirler: şifreleme, log kayıtları, erişim kontrolü (KVKK 12)",
        "kategori": "teknik",
        "oncelik": "yuksek",
    },
    {
        "no": 9,
        "madde": "İdari tedbirler: gizlilik sözleşmeleri, periyodik denetim (KVKK 12)",
        "kategori": "idari",
        "oncelik": "yuksek",
    },
    {
        "no": 10,
        "madde": "Yurt dışı veri aktarımı varsa açık rıza / yeterlilik kararı (KVKK 9)",
        "kategori": "transfer",
        "oncelik": "yuksek",
    },
    {
        "no": 11,
        "madde": "İlgili kişi başvurularını cevaplama prosedürü var mı? (KVKK 13 — 30 gün)",
        "kategori": "haklar",
        "oncelik": "yuksek",
    },
    {
        "no": 12,
        "madde": "Web sitesinde gizlilik politikası ve çerez aydınlatma metni yayında mı?",
        "kategori": "bilgilendirme",
        "oncelik": "orta",
    },
    {
        "no": 13,
        "madde": "Üçüncü taraf veri işleyenlerle KVKK uyumlu sözleşme imzalandı mı? (Yön. Md. 12)",
        "kategori": "idari",
        "oncelik": "yuksek",
    },
    {
        "no": 14,
        "madde": "Veri minimizasyonu uygulanıyor mu? (KVKK 4/d — ölçülülük)",
        "kategori": "envanter",
        "oncelik": "orta",
    },
    {
        "no": 15,
        "madde": "Veri envanterinde her veri kalemi için hukuki sebep belirtildi mi? (KVKK 5)",
        "kategori": "envanter",
        "oncelik": "yuksek",
    },
    {
        "no": 16,
        "madde": "İmha periyotları için tutanak / log saklanıyor mu? (Yön. Md. 11)",
        "kategori": "saklama",
        "oncelik": "orta",
    },
    {
        "no": 17,
        "madde": "Periyodik imha süresi (en geç 6 ayda bir) belirlendi mi? (Yön. Md. 11)",
        "kategori": "saklama",
        "oncelik": "orta",
    },
    {
        "no": 18,
        "madde": "Kişisel veri ihlali kayıt defteri tutuluyor mu?",
        "kategori": "ihlal",
        "oncelik": "orta",
    },
    {
        "no": 19,
        "madde": "Verileri yedekleme ve felaket kurtarma planı mevcut mu?",
        "kategori": "teknik",
        "oncelik": "orta",
    },
    {
        "no": 20,
        "madde": "Erişim yetkileri rol bazlı (RBAC) tanımlanmış mı?",
        "kategori": "teknik",
        "oncelik": "orta",
    },
    {
        "no": 21,
        "madde": "İlgili kişi başvuru formu (web/e-posta) hazır mı? (KVKK 13)",
        "kategori": "haklar",
        "oncelik": "orta",
    },
    {
        "no": 22,
        "madde": "Açık rıza ile gerekli olmayan hallerin ayrımı yapıldı mı? (KVKK 5/2)",
        "kategori": "rıza",
        "oncelik": "orta",
    },
    {
        "no": 23,
        "madde": "Çerez (cookie) tercih yönetimi (opt-in/opt-out) uygulanıyor mu?",
        "kategori": "bilgilendirme",
        "oncelik": "dusuk",
    },
]


# --------------------------------------------------------------------------- #
# Sektöre özel ek maddeler
# --------------------------------------------------------------------------- #

SEKTOR_EK_MADDELER: Dict[str, List[Dict[str, object]]] = {
    "saglik": [
        {"madde": "Özel nitelikli sağlık verisinde ek koruma tedbirleri (KVKK 6)", "kategori": "ozel_nitelikli", "oncelik": "yuksek"},
        {"madde": "Sağlık verisi için ayrı erişim logu ve şifreleme zorunlu (Kurul kararları)", "kategori": "teknik", "oncelik": "yuksek"},
        {"madde": "Sağlık verisi imha süresi hekim hizmeti bitiminden itibaren tıbbi mevzuata uygun (asgari 20 yıl)", "kategori": "saklama", "oncelik": "yuksek"},
        {"madde": "Hasta hakları yönetmeliği ile KVKK arasında uyum sağlandı mı?", "kategori": "sektorel", "oncelik": "yuksek"},
        {"madde": "MEDULA / e-Nabız entegrasyonlarında veri akış belgesi var mı?", "kategori": "sektorel", "oncelik": "orta"},
        {"madde": "Tetkik laboratuvarı / dış sağlık tedarikçilerle veri işleme sözleşmesi", "kategori": "idari", "oncelik": "yuksek"},
        {"madde": "Sağlık çalışanları için ek gizlilik yükümlülüğü ve sır saklama eğitimi", "kategori": "eğitim", "oncelik": "yuksek"},
    ],
    "fintech": [
        {"madde": "BDDK / MASAK / SPK uyumluluğu KVKK ile birlikte değerlendirildi mi?", "kategori": "sektorel", "oncelik": "yuksek"},
        {"madde": "Müşteri tanıma (KYC) verilerinin saklama süresi MASAK 8 yıl ile uyumlu mu?", "kategori": "saklama", "oncelik": "yuksek"},
        {"madde": "Kart verisi (PAN, CVV) için PCI-DSS uyumu sağlandı mı?", "kategori": "teknik", "oncelik": "yuksek"},
        {"madde": "Bankacılık sırrı / müşteri sırrı kapsamı (5411 sayılı Kanun 73) gözetildi mi?", "kategori": "sektorel", "oncelik": "yuksek"},
        {"madde": "Skoring / kredi notu işlemlerinde otomatik karar (KVKK 11/g) bilgilendirmesi", "kategori": "haklar", "oncelik": "orta"},
        {"madde": "Dolandırıcılık önleme (anti-fraud) sistemleri için meşru menfaat dengelemesi yapıldı mı?", "kategori": "rıza", "oncelik": "orta"},
    ],
    "egitim": [
        {"madde": "18 yaş altı öğrenciler için ebeveyn / veli rızası alınıyor mu?", "kategori": "rıza", "oncelik": "yuksek"},
        {"madde": "Akademik kayıt saklama süresi mevzuata uygun (mezuniyet sonrası saklama)", "kategori": "saklama", "oncelik": "yuksek"},
        {"madde": "MEB / YÖK veri paylaşım protokolleri belgelendirildi mi?", "kategori": "sektorel", "oncelik": "orta"},
        {"madde": "Uzaktan eğitim platformlarındaki kamera/ses kayıtları için ayrı rıza", "kategori": "rıza", "oncelik": "yuksek"},
        {"madde": "Öğrenci sağlık ve disiplin verileri için özel nitelikli koruma", "kategori": "ozel_nitelikli", "oncelik": "yuksek"},
        {"madde": "Sınav / değerlendirme verilerinin profilleme amacıyla işlenmesi şeffaf mı?", "kategori": "bilgilendirme", "oncelik": "orta"},
    ],
    "e_ticaret": [
        {"madde": "Tüketici verisi koruma (6502 sayılı Kanun) ile KVKK çakışmaları gözetildi mi?", "kategori": "sektorel", "oncelik": "yuksek"},
        {"madde": "Çerez politikası ayrıntılı yayında ve e-Privacy yaklaşımına uygun mu?", "kategori": "bilgilendirme", "oncelik": "yuksek"},
        {"madde": "Ticari elektronik ileti (İYS) onayı KVKK rızasından ayrı yönetiliyor mu?", "kategori": "rıza", "oncelik": "yuksek"},
        {"madde": "Sipariş / fatura verisi VUK 5 yıl, ticari belge 10 yıl saklama göz önünde mi?", "kategori": "saklama", "oncelik": "yuksek"},
        {"madde": "Ödeme / kart verisi PCI-DSS ve tokenizasyon ile korunuyor mu?", "kategori": "teknik", "oncelik": "yuksek"},
        {"madde": "Kargo / lojistik tedarikçileri için veri işleyen sözleşmesi var mı?", "kategori": "idari", "oncelik": "orta"},
        {"madde": "Davranışsal reklam (retargeting) için ayrı rıza alınıyor mu?", "kategori": "rıza", "oncelik": "orta"},
    ],
    "imalat": [
        {"madde": "Tedarikçi / bayi personel verisi için işleme sözleşmesi", "kategori": "idari", "oncelik": "orta"},
        {"madde": "Üretim süreçlerinde işlenen veri (telemetri / IoT) anonimleştirme politikası", "kategori": "teknik", "oncelik": "orta"},
        {"madde": "Fabrika CCTV kayıtları için aydınlatma levhası ve saklama süresi (asgari 30 gün)", "kategori": "sektorel", "oncelik": "orta"},
        {"madde": "OHS (İSG) kapsamında sağlık taramaları için özel nitelikli veri prosedürü", "kategori": "ozel_nitelikli", "oncelik": "yuksek"},
        {"madde": "Ar-Ge bilgilerinde ticari sır / kişisel veri ayrımı yapıldı mı?", "kategori": "envanter", "oncelik": "orta"},
    ],
    "kamu": [
        {"madde": "Kamu otoritelerine veri paylaşımının hukuki sebebi (KVKK 8/2) belgelendi mi?", "kategori": "transfer", "oncelik": "yuksek"},
        {"madde": "CİMER / BİMER başvurularında ilgili kişi haklarıyla çakışma yönetimi", "kategori": "haklar", "oncelik": "orta"},
        {"madde": "Kamu görevlileri için ayrı veri sorumlusu / yetkili tanımı yapıldı mı?", "kategori": "yönetim", "oncelik": "yuksek"},
        {"madde": "5018 sayılı Kanun ve KVKK uyumu kapsamında denetim raporları", "kategori": "sektorel", "oncelik": "orta"},
        {"madde": "Açık veri / şeffaflık yükümlülüğü ile kişisel veri koruma dengesi", "kategori": "bilgilendirme", "oncelik": "orta"},
    ],
    "telekom": [
        {"madde": "BTK Elektronik Haberleşme Sektöründe Kişisel Verilerin İşlenmesi Yönetmeliği uyumu", "kategori": "sektorel", "oncelik": "yuksek"},
        {"madde": "Trafik / konum verisi saklama süresi 2 yıl ile sınırlı mı? (Yön. Md. 9)", "kategori": "saklama", "oncelik": "yuksek"},
        {"madde": "Abone iletişim verisinin yargı taleplerine cevap prosedürü tanımlı mı?", "kategori": "transfer", "oncelik": "yuksek"},
        {"madde": "İletişim kaydı için ek aydınlatma ve gizliliği koruma tedbirleri", "kategori": "bilgilendirme", "oncelik": "yuksek"},
    ],
    "insan_kaynaklari": [
        {"madde": "Aday başvuru CV'lerinin değerlendirme sonrası saklama süresi (azami 2 yıl önerilir)", "kategori": "saklama", "oncelik": "yuksek"},
        {"madde": "Çalışan referans / soruşturma için açık rıza alındı mı?", "kategori": "rıza", "oncelik": "yuksek"},
        {"madde": "Performans değerlendirme verileri için aydınlatma yapıldı mı?", "kategori": "bilgilendirme", "oncelik": "orta"},
        {"madde": "İşten ayrılan personel verisinin SGK / İş K. açısından saklama süresi (10 yıl)", "kategori": "saklama", "oncelik": "yuksek"},
        {"madde": "Çalışan e-posta / cihaz izleme için meşru menfaat dengelemesi yapıldı mı?", "kategori": "idari", "oncelik": "yuksek"},
    ],
    "hukuk_burosu": [
        {"madde": "Müvekkil sırrı (Av.K. 36) ve KVKK uyumu için ayrı politika", "kategori": "sektorel", "oncelik": "yuksek"},
        {"madde": "Dava dosyalarının saklama süresi avukatlık mevzuatına uygun (asgari 10 yıl)", "kategori": "saklama", "oncelik": "yuksek"},
        {"madde": "Karşı taraf verisinin işlenmesinde hukuki sebep belirlendi mi? (KVKK 5/2-e)", "kategori": "rıza", "oncelik": "orta"},
        {"madde": "Dijital UYAP entegrasyonlarında veri akış belgesi", "kategori": "transfer", "oncelik": "orta"},
    ],
    "diger": [],
}


# --------------------------------------------------------------------------- #
# Veri türüne özel ek maddeler
# --------------------------------------------------------------------------- #

VERI_TURU_EK_MADDELER: Dict[str, List[Dict[str, object]]] = {
    "ozel_nitelikli": [
        {"madde": "Özel nitelikli veri için Kurul'un belirlediği yeterli önlemler alındı mı? (KVKK 6/4)", "kategori": "ozel_nitelikli", "oncelik": "yuksek"},
        {"madde": "Özel nitelikli veri için ayrı politika ve prosedür dokümanı var mı?", "kategori": "ozel_nitelikli", "oncelik": "yuksek"},
    ],
    "saglik": [
        {"madde": "Sağlık verisi işleme yalnızca sır saklama yükümlülüğü altındaki kişilerce yapılıyor mu? (KVKK 6/3)", "kategori": "ozel_nitelikli", "oncelik": "yuksek"},
    ],
    "cocuk": [
        {"madde": "18 yaş altı veri işlemede veli/vasi rızası ve yaş doğrulama mekanizması var mı?", "kategori": "rıza", "oncelik": "yuksek"},
    ],
    "biyometrik": [
        {"madde": "Biyometrik veri için yedek (biyometrik olmayan) alternatif sunuluyor mu?", "kategori": "ozel_nitelikli", "oncelik": "yuksek"},
        {"madde": "Biyometrik şablonların ham görüntü yerine matematiksel temsille saklanması", "kategori": "teknik", "oncelik": "yuksek"},
    ],
    "kamera": [
        {"madde": "CCTV alanında aydınlatma levhası asıldı mı? (KVKK Kurul 2018/63 kararı)", "kategori": "bilgilendirme", "oncelik": "yuksek"},
        {"madde": "CCTV kayıt süresi ihtiyaca uygun ölçülü mü? (asgari, makul süre)", "kategori": "saklama", "oncelik": "orta"},
    ],
    "konum": [
        {"madde": "Konum verisi için ayrı açık rıza ve uygulama içinde anlık kontrol seçeneği", "kategori": "rıza", "oncelik": "yuksek"},
    ],
    "iletisim_kaydi": [
        {"madde": "Çağrı kaydı için aydınlatma anonsu yapılıyor mu?", "kategori": "bilgilendirme", "oncelik": "yuksek"},
        {"madde": "Çağrı kayıtları için saklama süresi makul ve sebebe bağlı mı?", "kategori": "saklama", "oncelik": "orta"},
    ],
    "finansal": [
        {"madde": "Finansal veri için ek şifreleme, maskleme ve erişim kontrolü uygulandı mı?", "kategori": "teknik", "oncelik": "yuksek"},
    ],
    "cerez": [
        {"madde": "Zorunlu olmayan çerezler için açık rıza (banner / consent management)", "kategori": "rıza", "oncelik": "yuksek"},
    ],
    "calisan": [
        {"madde": "İK verisi için ayrı aydınlatma metni hazırlandı mı?", "kategori": "bilgilendirme", "oncelik": "orta"},
    ],
    "musteri": [],
    "kisisel": [],
}


YASAL_UYARI = (
    "Bu rehber, profesyonel KVKK danışmanlığının yerine geçmez. KVKK uyumluluğu "
    "her kurum için özgün koşullara göre değerlendirilmelidir. Çıktılar bir AI "
    "taslağıdır; sertifikalı bir KVKK uzmanı veya hukuk müşaviri tarafından "
    "incelenmelidir."
)


# --------------------------------------------------------------------------- #
# Yardımcı fonksiyonlar
# --------------------------------------------------------------------------- #

def list_sektorler() -> List[str]:
    """Mevcut sektör anahtarlarını döndürür."""
    return list(SEKTOR_ETIKETLERI.keys())


def list_veri_turleri() -> List[str]:
    """Mevcut veri türü anahtarlarını döndürür."""
    return list(VERI_TURU_ETIKETLERI.keys())


def sektor_etiketi(sektor: str) -> str:
    """UI için okunabilir sektör etiketi."""
    return SEKTOR_ETIKETLERI.get(sektor, sektor)


def veri_turu_etiketi(veri_turu: str) -> str:
    """UI için okunabilir veri türü etiketi."""
    return VERI_TURU_ETIKETLERI.get(veri_turu, veri_turu)


def kategori_etiketi(kategori: str) -> str:
    """UI için okunabilir kategori etiketi."""
    return KATEGORI_ETIKETLERI.get(kategori, kategori.capitalize())


def _normalize_oncelik(deger: Optional[str]) -> str:
    """Önceliği yuksek/orta/dusuk değerlerinden birine indirger."""
    if not deger:
        return "orta"
    d = str(deger).strip().lower()
    if d in ("yuksek", "yüksek", "high", "critical", "kritik"):
        return "yuksek"
    if d in ("dusuk", "düşük", "low"):
        return "dusuk"
    return "orta"


def _safe_json_extract(text: str) -> Optional[dict]:
    """LLM çıktısından ilk JSON objesini parse eder."""
    if not text:
        return None
    # ```json ... ``` blokları temizle
    fenced = re.search(r"```(?:json)?\s*(.*?)```", text, re.DOTALL)
    candidate = fenced.group(1) if fenced else text
    # İlk { ... } bloğunu yakala
    m = re.search(r"\{.*\}", candidate, re.DOTALL)
    if not m:
        return None
    try:
        return json.loads(m.group(0))
    except json.JSONDecodeError:
        return None


def _llm_ek_uret(sektor: str, veri_turleri: List[str]) -> dict:
    """LLM ile sektöre özel ek maddeler ve notlar üretir.

    Returns:
        {
            "ek_maddeler": [{"madde": str, "kategori": str, "oncelik": str}],
            "sektorel_notlar": {kategori: str},
            "ozet": str,
        }
        Hata halinde boş yapı döner.
    """
    bos = {"ek_maddeler": [], "sektorel_notlar": {}, "ozet": ""}

    if llm_generate is None:
        return bos
    try:
        if llm_available is not None and not llm_available():
            return bos
    except Exception:
        return bos

    sektor_label = sektor_etiketi(sektor)
    veri_labels = [veri_turu_etiketi(v) for v in veri_turleri]

    system = (
        "Sen Türkiye'de 6698 sayılı KVKK ve ilgili ikincil mevzuatta uzman "
        "bir hukuk / uyum danışmanısın. Çıktıyı yalnızca geçerli JSON "
        "formatında ver, açıklama yazma."
    )
    user = (
        f"Sektör: {sektor_label}\n"
        f"İşlenen veri türleri: {', '.join(veri_labels) if veri_labels else '(belirtilmedi)'}\n\n"
        "Aşağıdaki yapıda, BU SEKTÖRE özel KVKK uyum ek maddelerini ve "
        "sektörel notları üret. Maddeler, genel KVKK gerekliliklerini "
        "TEKRAR ETMEMELİ; yalnızca sektöre / veri türüne özgü ek konulara "
        "odaklanmalı. 5-10 madde yeterlidir. Kategoriler şu setten seçilmeli: "
        "[bilgilendirme, rıza, envanter, saklama, ihlal, yönetim, eğitim, "
        "teknik, idari, transfer, haklar, ozel_nitelikli, sektorel]. "
        "Öncelik: yuksek, orta veya dusuk olmalı.\n\n"
        "Format (sadece JSON):\n"
        "{\n"
        '  "ek_maddeler": [\n'
        '    {"madde": "...", "kategori": "...", "oncelik": "yuksek|orta|dusuk"}\n'
        "  ],\n"
        '  "sektorel_notlar": {"teknik": "...", "saklama": "..."},\n'
        '  "ozet": "Bu sektör için en kritik 1-2 cümlelik genel değerlendirme."\n'
        "}\n"
    )

    try:
        raw = llm_generate(system=system, user=user, max_tokens=1500, temperature=0.2)
    except Exception:
        return bos

    parsed = _safe_json_extract(raw)
    if not parsed:
        return bos

    ek_maddeler: List[Dict[str, object]] = []
    for item in parsed.get("ek_maddeler", []) or []:
        if not isinstance(item, dict):
            continue
        metin = str(item.get("madde", "")).strip()
        if not metin:
            continue
        kategori = str(item.get("kategori", "sektorel")).strip().lower()
        if kategori not in KATEGORI_ETIKETLERI:
            kategori = "sektorel"
        ek_maddeler.append(
            {
                "madde": metin,
                "kategori": kategori,
                "oncelik": _normalize_oncelik(item.get("oncelik")),
            }
        )

    sektorel_notlar = parsed.get("sektorel_notlar") or {}
    if not isinstance(sektorel_notlar, dict):
        sektorel_notlar = {}
    # Yalnızca bilinen kategorileri tut
    sektorel_notlar = {
        str(k): str(v)
        for k, v in sektorel_notlar.items()
        if str(k) in KATEGORI_ETIKETLERI and v
    }

    return {
        "ek_maddeler": ek_maddeler,
        "sektorel_notlar": sektorel_notlar,
        "ozet": str(parsed.get("ozet", "")).strip(),
    }


def _uyum_skoru_tahmin(maddeler: List[Dict[str, object]]) -> int:
    """Toplam madde sayısı ve önceliklere bakarak başlangıç (referans) skoru tahmin et.

    Bu, kullanıcı hiçbir kutuyu işaretlemeden önceki "ortalama olgunluk"
    referansıdır. Statik bir tahmin: yüksek öncelikli madde sayısı arttıkça
    referans 50-65 arasında değişir.
    """
    if not maddeler:
        return 0
    yuksek = sum(1 for m in maddeler if m.get("oncelik") == "yuksek")
    orta = sum(1 for m in maddeler if m.get("oncelik") == "orta")
    dusuk = len(maddeler) - yuksek - orta
    # Referans skor: maksimum 60 olsun (kullanıcı uyum sağladıkça artar)
    payda = 3 * yuksek + 2 * orta + 1 * dusuk
    if payda == 0:
        return 50
    # Yüksek payda yoğunluğu -> daha düşük referans (daha çok iş var)
    skor = max(40, 70 - int(payda * 1.2))
    return min(skor, 70)


def uyum_skoru_hesapla(maddeler: List[Dict[str, object]], tamamlananlar: List[int]) -> int:
    """Kullanıcının işaretlediği maddelere göre 0-100 ağırlıklı uyum skoru.

    Ağırlıklar:
        yuksek = 3, orta = 2, dusuk = 1
    """
    if not maddeler:
        return 0
    agirlik_map = {"yuksek": 3, "orta": 2, "dusuk": 1}
    toplam = 0
    kazanilan = 0
    tamam_set = set(tamamlananlar or [])
    for m in maddeler:
        w = agirlik_map.get(str(m.get("oncelik", "orta")), 2)
        toplam += w
        if int(m.get("no", -1)) in tamam_set:
            kazanilan += w
    if toplam == 0:
        return 0
    return round(100 * kazanilan / toplam)


# --------------------------------------------------------------------------- #
# Ana fonksiyon
# --------------------------------------------------------------------------- #

def checklist_uret(
    sektor: str,
    veri_turleri: List[str],
    llm_ek: bool = True,
) -> dict:
    """Sektör ve veri türlerine göre KVKK uyum checklist'i üretir.

    Args:
        sektor: SEKTOR_ETIKETLERI anahtarlarından biri.
        veri_turleri: VERI_TURU_ETIKETLERI anahtarlarından oluşan liste.
        llm_ek: True ise LLM ile sektöre özel ek maddeler üretilir.

    Returns:
        {
            "maddeler": [
                {
                    "no": int,
                    "madde": str,
                    "kategori": str,
                    "oncelik": "yuksek"|"orta"|"dusuk",
                    "sektorel_not": str,
                },
                ...
            ],
            "sektor": str,
            "sektor_label": str,
            "veri_turleri": [str],
            "ozet": str,
            "tahmin_uyum_skoru": int (0-100),
            "yasal_uyari": str,
            "llm_kullanildi": bool,
        }
    """
    sektor = (sektor or "diger").lower()
    if sektor not in SEKTOR_ETIKETLERI:
        sektor = "diger"
    veri_turleri = list(veri_turleri or [])

    # 1) Temel maddeler
    maddeler: List[Dict[str, object]] = [dict(m) for m in TEMEL_CHECKLIST]

    # 2) Sektörel maddeler
    sektor_maddeleri = SEKTOR_EK_MADDELER.get(sektor, []) or []

    # 3) Veri türü maddeleri (tekrarsız ekleme için "madde" metnine göre)
    veri_maddeleri: List[Dict[str, object]] = []
    for vt in veri_turleri:
        for em in VERI_TURU_EK_MADDELER.get(vt, []) or []:
            veri_maddeleri.append(em)

    # 4) LLM ek maddeleri
    llm_kullanildi = False
    llm_ek_sonuc: dict = {"ek_maddeler": [], "sektorel_notlar": {}, "ozet": ""}
    if llm_ek:
        llm_ek_sonuc = _llm_ek_uret(sektor, veri_turleri)
        if llm_ek_sonuc.get("ek_maddeler"):
            llm_kullanildi = True

    # Tüm ek maddeleri birleştir, metin bazında deduplicate et
    gorulen_metinler = {str(m["madde"]).lower() for m in maddeler}

    def _ekle(ek_list: List[Dict[str, object]]) -> None:
        nonlocal maddeler, gorulen_metinler
        for em in ek_list:
            metin = str(em.get("madde", "")).strip()
            if not metin:
                continue
            anahtar = metin.lower()
            if anahtar in gorulen_metinler:
                continue
            gorulen_metinler.add(anahtar)
            no = len(maddeler) + 1
            maddeler.append(
                {
                    "no": no,
                    "madde": metin,
                    "kategori": str(em.get("kategori", "sektorel")),
                    "oncelik": _normalize_oncelik(em.get("oncelik")),
                }
            )

    _ekle(sektor_maddeleri)
    _ekle(veri_maddeleri)
    _ekle(llm_ek_sonuc.get("ek_maddeler", []))

    # 5) Sektörel notları her maddeye yaz (LLM'den gelirse)
    sektorel_notlar: Dict[str, str] = llm_ek_sonuc.get("sektorel_notlar", {}) or {}
    for m in maddeler:
        not_metni = sektorel_notlar.get(str(m["kategori"]), "")
        m["sektorel_not"] = not_metni

    # 6) Özet
    sektor_label = sektor_etiketi(sektor)
    veri_label = ", ".join(veri_turu_etiketi(v) for v in veri_turleri) or "(belirtilmedi)"
    llm_ozet = llm_ek_sonuc.get("ozet", "") or ""
    ozet_satirlari = [
        f"{sektor_label} sektörü için {len(maddeler)} maddelik KVKK uyum checklist'i üretildi.",
        f"İşlenen veri türleri: {veri_label}.",
    ]
    if llm_ozet:
        ozet_satirlari.append(llm_ozet)
    ozet = " ".join(ozet_satirlari)

    # 7) Tahmin skor
    tahmin = _uyum_skoru_tahmin(maddeler)

    return {
        "maddeler": maddeler,
        "sektor": sektor,
        "sektor_label": sektor_label,
        "veri_turleri": veri_turleri,
        "ozet": ozet,
        "tahmin_uyum_skoru": tahmin,
        "yasal_uyari": YASAL_UYARI,
        "llm_kullanildi": llm_kullanildi,
    }


# --------------------------------------------------------------------------- #
# Dışa aktarım yardımcıları
# --------------------------------------------------------------------------- #

def maddeleri_kategoriye_grupla(
    maddeler: List[Dict[str, object]],
) -> Dict[str, List[Dict[str, object]]]:
    """Maddeleri kategoriye göre grupla. Kategorilerin sırası KATEGORI_ETIKETLERI sırasıdır."""
    gruplu: Dict[str, List[Dict[str, object]]] = {k: [] for k in KATEGORI_ETIKETLERI}
    for m in maddeler:
        k = str(m.get("kategori", "sektorel"))
        gruplu.setdefault(k, []).append(m)
    # Boş kategorileri at
    return {k: v for k, v in gruplu.items() if v}


def to_markdown(sonuc: dict, tamamlananlar: Optional[List[int]] = None) -> str:
    """Checklist sonucunu Markdown'a dönüştür."""
    tamam_set = set(tamamlananlar or [])
    maddeler = sonuc.get("maddeler", []) or []
    sektor_label = sonuc.get("sektor_label", "")
    veri_turleri = sonuc.get("veri_turleri", []) or []
    veri_label = ", ".join(veri_turu_etiketi(v) for v in veri_turleri) or "—"

    skor = uyum_skoru_hesapla(maddeler, list(tamam_set))

    lines: List[str] = []
    lines.append("# KVKK Uyum Kontrol Listesi")
    lines.append("")
    lines.append(f"**Sektör:** {sektor_label}")
    lines.append(f"**İşlenen veri türleri:** {veri_label}")
    lines.append(f"**Toplam madde:** {len(maddeler)}")
    lines.append(f"**Uyum Skoru:** {skor}/100")
    lines.append("")
    if sonuc.get("ozet"):
        lines.append("## Özet")
        lines.append(str(sonuc["ozet"]))
        lines.append("")

    gruplu = maddeleri_kategoriye_grupla(maddeler)
    for kategori, items in gruplu.items():
        lines.append(f"## {kategori_etiketi(kategori)}")
        lines.append("")
        for m in items:
            no = m.get("no", "?")
            tik = "[x]" if int(no) in tamam_set else "[ ]"
            oncelik = str(m.get("oncelik", "orta")).upper()
            lines.append(f"- {tik} **#{no}** ({oncelik}) {m.get('madde','')}")
            not_metni = str(m.get("sektorel_not", "")).strip()
            if not_metni:
                lines.append(f"    - _Sektörel not:_ {not_metni}")
        lines.append("")

    lines.append("---")
    lines.append("")
    lines.append("## Yasal Uyarı")
    lines.append(sonuc.get("yasal_uyari", YASAL_UYARI))
    return "\n".join(lines)


def to_docx_bytes(sonuc: dict, tamamlananlar: Optional[List[int]] = None) -> bytes:
    """Checklist sonucunu .docx olarak byte halinde döndür."""
    from io import BytesIO

    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as e:
        raise RuntimeError(
            "python-docx yüklü değil. `pip install python-docx` ile kurun."
        ) from e

    tamam_set = set(tamamlananlar or [])
    maddeler = sonuc.get("maddeler", []) or []
    sektor_label = sonuc.get("sektor_label", "")
    veri_turleri = sonuc.get("veri_turleri", []) or []
    veri_label = ", ".join(veri_turu_etiketi(v) for v in veri_turleri) or "—"
    skor = uyum_skoru_hesapla(maddeler, list(tamam_set))

    doc = Document()
    doc.add_heading("KVKK Uyum Kontrol Listesi", level=0)

    p = doc.add_paragraph()
    p.add_run("Sektör: ").bold = True
    p.add_run(str(sektor_label))

    p = doc.add_paragraph()
    p.add_run("İşlenen veri türleri: ").bold = True
    p.add_run(veri_label)

    p = doc.add_paragraph()
    p.add_run("Toplam madde: ").bold = True
    p.add_run(str(len(maddeler)))

    p = doc.add_paragraph()
    p.add_run("Uyum Skoru: ").bold = True
    p.add_run(f"{skor}/100")

    if sonuc.get("ozet"):
        doc.add_heading("Özet", level=1)
        doc.add_paragraph(str(sonuc["ozet"]))

    gruplu = maddeleri_kategoriye_grupla(maddeler)
    for kategori, items in gruplu.items():
        doc.add_heading(kategori_etiketi(kategori), level=1)
        for m in items:
            no = m.get("no", "?")
            tik = "[X]" if int(no) in tamam_set else "[ ]"
            oncelik = str(m.get("oncelik", "orta")).upper()
            metin = f"{tik}  #{no} ({oncelik})  {m.get('madde','')}"
            para = doc.add_paragraph(metin, style="List Bullet")
            for run in para.runs:
                run.font.size = Pt(10)
            not_metni = str(m.get("sektorel_not", "")).strip()
            if not_metni:
                sub = doc.add_paragraph(f"Sektörel not: {not_metni}")
                sub.paragraph_format.left_indent = Pt(24)
                for run in sub.runs:
                    run.italic = True
                    run.font.size = Pt(9)

    doc.add_heading("Yasal Uyarı", level=1)
    doc.add_paragraph(sonuc.get("yasal_uyari", YASAL_UYARI))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_pdf_bytes(sonuc: dict, tamamlananlar: Optional[List[int]] = None) -> bytes:
    """Checklist sonucunu .pdf olarak byte halinde döndür.

    reportlab varsa onu kullanır, yoksa fpdf2'ye düşer.
    """
    tamam_set = set(tamamlananlar or [])
    maddeler = sonuc.get("maddeler", []) or []
    sektor_label = sonuc.get("sektor_label", "")
    veri_turleri = sonuc.get("veri_turleri", []) or []
    veri_label = ", ".join(veri_turu_etiketi(v) for v in veri_turleri) or "—"
    skor = uyum_skoru_hesapla(maddeler, list(tamam_set))

    # --- reportlab yolu (Türkçe karakter desteği için Helvetica yeterli) ---
    try:
        from io import BytesIO
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer,
            PageBreak,
        )

        buf = BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            leftMargin=2 * cm,
            rightMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
            title="KVKK Uyum Kontrol Listesi",
        )
        styles = getSampleStyleSheet()
        h_style = styles["Heading1"]
        h2_style = styles["Heading2"]
        body = styles["BodyText"]
        small = ParagraphStyle(
            "small", parent=body, fontSize=8, textColor="#666666", leftIndent=10
        )

        elements = []
        elements.append(Paragraph("KVKK Uyum Kontrol Listesi", h_style))
        elements.append(Spacer(1, 6))
        elements.append(Paragraph(f"<b>Sektör:</b> {sektor_label}", body))
        elements.append(Paragraph(f"<b>Veri türleri:</b> {veri_label}", body))
        elements.append(Paragraph(f"<b>Toplam madde:</b> {len(maddeler)}", body))
        elements.append(Paragraph(f"<b>Uyum Skoru:</b> {skor}/100", body))
        elements.append(Spacer(1, 12))

        if sonuc.get("ozet"):
            elements.append(Paragraph("Özet", h2_style))
            elements.append(Paragraph(str(sonuc["ozet"]), body))
            elements.append(Spacer(1, 8))

        gruplu = maddeleri_kategoriye_grupla(maddeler)
        for kategori, items in gruplu.items():
            elements.append(Paragraph(kategori_etiketi(kategori), h2_style))
            for m in items:
                no = m.get("no", "?")
                tik = "[X]" if int(no) in tamam_set else "[ ]"
                oncelik = str(m.get("oncelik", "orta")).upper()
                metin = (
                    f"{tik} <b>#{no}</b> ({oncelik}) "
                    f"{str(m.get('madde','')).replace('<','&lt;').replace('>','&gt;')}"
                )
                elements.append(Paragraph(metin, body))
                not_metni = str(m.get("sektorel_not", "")).strip()
                if not_metni:
                    elements.append(
                        Paragraph(
                            f"<i>Sektörel not:</i> "
                            f"{not_metni.replace('<','&lt;').replace('>','&gt;')}",
                            small,
                        )
                    )
            elements.append(Spacer(1, 6))

        elements.append(PageBreak())
        elements.append(Paragraph("Yasal Uyarı", h2_style))
        elements.append(Paragraph(sonuc.get("yasal_uyari", YASAL_UYARI), body))

        doc.build(elements)
        return buf.getvalue()
    except ImportError:
        pass

    # --- fpdf2 yolu ---
    try:
        from io import BytesIO
        from fpdf import FPDF

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, "KVKK Uyum Kontrol Listesi", ln=1)

        pdf.set_font("Helvetica", "", 11)
        pdf.cell(0, 7, f"Sektor: {sektor_label}", ln=1)
        pdf.cell(0, 7, f"Veri turleri: {veri_label}", ln=1)
        pdf.cell(0, 7, f"Toplam madde: {len(maddeler)}", ln=1)
        pdf.cell(0, 7, f"Uyum Skoru: {skor}/100", ln=1)
        pdf.ln(4)

        if sonuc.get("ozet"):
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, "Ozet", ln=1)
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(0, 6, str(sonuc["ozet"]))
            pdf.ln(2)

        gruplu = maddeleri_kategoriye_grupla(maddeler)
        for kategori, items in gruplu.items():
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, kategori_etiketi(kategori), ln=1)
            pdf.set_font("Helvetica", "", 10)
            for m in items:
                no = m.get("no", "?")
                tik = "[X]" if int(no) in tamam_set else "[ ]"
                oncelik = str(m.get("oncelik", "orta")).upper()
                pdf.multi_cell(
                    0, 6, f"{tik} #{no} ({oncelik})  {m.get('madde','')}"
                )
                not_metni = str(m.get("sektorel_not", "")).strip()
                if not_metni:
                    pdf.set_font("Helvetica", "I", 9)
                    pdf.multi_cell(0, 5, f"    Sektorel not: {not_metni}")
                    pdf.set_font("Helvetica", "", 10)
            pdf.ln(1)

        pdf.add_page()
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, "Yasal Uyari", ln=1)
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 6, sonuc.get("yasal_uyari", YASAL_UYARI))

        out = pdf.output(dest="S")
        if isinstance(out, str):
            return out.encode("latin-1", errors="replace")
        return bytes(out)
    except ImportError as e:
        raise RuntimeError(
            "PDF üretimi için `reportlab` veya `fpdf2` paketlerinden biri gereklidir."
        ) from e
